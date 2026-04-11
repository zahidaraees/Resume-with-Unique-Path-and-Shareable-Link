from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Name
name = doc.add_paragraph('Zahida Raees')
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name.runs[0].font.size = Pt(24)
name.runs[0].bold = True

# Contact Info
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.add_run('Cell#: +92-333-2203353\n')
contact.add_run('Email: zahidaraeesi@gmail.com\n')
contact.add_run('Ittifaq Mention, 1st Fl, Flat #. 02 Plot No LY 14/172, St# 3, West Yousuf Haroon Road, Karachi 53, Post Code 75660, Karachi.')

# Objective
doc.add_heading('Objective', level=2)
doc.add_paragraph('I am always ready to work in any challenging environment to utilize my education, experience and professional skills.')

# Academic & Professional Education
doc.add_heading('Academic & Professional Education', level=2)
doc.add_paragraph('• Student of Cloud Applied Generative AI Engineering Q4, 1st batch, at governor house Karachi since Feb 2024')
doc.add_paragraph('• MPhil in Balochi language & Literature, University of Balochistan, Quetta (2019-2022)')
doc.add_paragraph('• Masters in Balochi Language and literature from the University of Balochistan, Quetta (2017)')
doc.add_paragraph('• One-year diploma in ORACLE 9i (DBA & DEVELOPER TRACK) from ORASOFT Karachi. May 2004 to April 2005')
doc.add_paragraph('• Six months short course Visual Basic 5 & SQL Server (2000), Karachi')
doc.add_paragraph('• Two Years Diploma in Computer Software, from Sindh Board of technical Education Karachi (1996-1998)')
doc.add_paragraph('• Bachelor of Science from Karachi University (1996)')

# Research Papers
doc.add_heading('Research Papers in Academic Research Journals', level=2)
doc.add_paragraph('1. "Balochistaniyat", 2022, "Mir Bibagar Rind", Published by Balochi Academy Quetta')
doc.add_paragraph('2. "Classical Poetry", annual Research Journal 2021, "Morphological Aspect of the selective Balochi Classical Poetry", published by Balochi Academy Quetta')
doc.add_paragraph('3. "Balocistaniyat", 2020, "History of Baloch & Balochi; the role of Orientalists", Published by Balochi Academy Quetta')

# Software Skills
doc.add_heading('Software Skills and Knowledge', level=2)
doc.add_paragraph('Full Stack Developer. Agentic AI Development: Claude, Gemini, Qwen, Docusaurus')
doc.add_paragraph('Web Applications and Website developments. Python, TypeScript, React, Tailwind CSS, Next-JS, Streamlit, Vercel, Github, Prompt Engineering, Context Engineering')
doc.add_paragraph('Visual Fox (all versions), CDS, Dream viewer')
doc.add_paragraph('ORACLE 9I, SQL, PL.SQL, DEVELOPER 6I, SQL SERVER 6.5, MSSQL')
doc.add_paragraph('Coral draw, Photoshop, DHTML, CSS, PHP, ASP, JSP, JQUERY, AJAX, CPANEL, Team viewer')
doc.add_paragraph('MS OFFICE, VISIO TECH, InPage (All versions), AUTOCAD, ULEAD, SONY VEGAS, Adobe Premier')
doc.add_paragraph('SMF, WORDPRESS')

# Literary Skills
doc.add_heading('Literary Skills', level=2)
doc.add_paragraph('• Poems and prose (short Stories, essays, novel, novelette, drama, book reviews and technical articles)')
doc.add_paragraph('• Translated literary and academic books, articles and research papers as well as own poetry and proses into Urdu')
doc.add_paragraph('• Wrote academic and non-academic research papers for the various journals, seminars and conferences')
doc.add_paragraph('• Online and onsite teaching experience of Balochi language')

# Published Books
doc.add_heading('Published Books', level=2)
doc.add_paragraph('• Book: سستگیں مُروارد - Research Paper along with collection of Balochi aphorisms (January 2020)')
doc.add_paragraph('• Book: توکہ زند ءِ دَروَرئے - Balochi Novel (June 2019)')
doc.add_paragraph('• Book: پادءِ چیر ءِ زمین - Balochi Novelette (July 2018)')
doc.add_paragraph('• Book: بے کسّہ ءَ من ءَ واب نئیت - Collection of Baloch Short Stories (2016)')
doc.add_paragraph('• Book: خواب گھروندا - Collection of Urdu Poetry (2014)')

# Literary Awards
doc.add_heading('Literary Awards', level=2)
doc.add_paragraph('• بلوچی ناول پہلی پوزیشن 2025. علامہ اقبال ادبی ایوارڈ۔ ادارہ ثقافت بلوچستان کوئٹہ')
doc.add_paragraph('• صبا داد 2023. سالانہ ادبی ایوارڈ منجانب سید ریفرنس لائبریری۔ کراچی')
doc.add_paragraph('• بلوچی لبزانک داد۔ 2018. منجانب ارشاد اسلامی۔ نیکشهر، سیستان و بلوچستان۔ ایران')
doc.add_paragraph('• بیگل داد۔ 2016۔ منجانب کویت ادبی انجمن۔ کویت')
doc.add_paragraph('• آزاد جمالدینی ایوارڈ۔ 2013۔ منجانب راسکو ادبی گل نوشکی، بلوچستان')

# Work Experience
doc.add_heading('Work Experience', level=2)

doc.add_paragraph('Teaching Assistant (Balochi Teacher)')
doc.add_paragraph('University of Karachi, Sindhi Department, Since Jan 2024')
doc.add_paragraph('I am teaching Balochi Minor Subject in BS First Year. First Semester (301), First year 2nd Semester (302), Second Year first Semester (401), Second Year Second Semester (402) as well as Balochi Certificate Course (one year), Diploma (two years) (All has 3 credit hours)')

doc.add_paragraph('Technical Work Experience: Since 2024')
doc.add_paragraph('• Physical AI and Humanoid Robotic Course application development and deployment @ github and vercel')
doc.add_paragraph('• Figma to next-js full stack e-commerce application development for furniture business')
doc.add_paragraph('• ZR-Portfolio')
doc.add_paragraph('• All rest of the projects are here: https://github.com/zahidaraees?tab=repositories')

doc.add_paragraph('Since 2005: Baask Communication Network (Non-profitable)')
doc.add_paragraph('• Development and administration of First Digital library of Balochi language and Literature named baask.com')
doc.add_paragraph('• Designed books, covers, cards, banners, brochures, pamphlets on demand of the offline and online customers')
doc.add_paragraph('• Created web pages, social media pages, groups and blogs for the different individuals and voluntary business groups')
doc.add_paragraph('• Consultancy on Online and offline applications')

doc.add_paragraph('21st May 2008 till Dec 2018: Software Consultant & Developer at Diversa Logic Pvt. Ltd.')
doc.add_paragraph('• Designed and support following business applications:')
doc.add_paragraph('  - Supply chain management (Premier Systems Private Limited)')
doc.add_paragraph('  - PPL application for Epson Products (ABM SERVICE Karachi)')
doc.add_paragraph('  - Sales and inventory control System (ABM SERVICE Karachi)')
doc.add_paragraph('  - Software Audit as an external auditor (Habib Insurance - Head office KARACHI)')
doc.add_paragraph('  - Billing System (GCG Billing System NY-USA)')
doc.add_paragraph('  - Sales Purchase system (American Fumigation Karachi)')

doc.add_paragraph('August 2000 to May 2008: Application Administrator & Developer at ABM Data Systems (Pvt.) Ltd')
doc.add_paragraph('• analysis, development, implementation and administration of financial accounting systems, call login systems, payroll systems, inventory control systems, maintenance contract and sales purchase systems of the all branches of the company (Karachi, Lahore, Islamabad and Hyderabad)')
doc.add_paragraph('• All kind of report generation from the application, on the request of upper and middle management')
doc.add_paragraph('• System applications auditing on quarterly, half yearly and yearly bases')
doc.add_paragraph('• Training and presentation of the application system to the individual users')
doc.add_paragraph('• Prompt response on any kind of query from any department of the company')
doc.add_paragraph('• Troubleshooting and support')

doc.add_paragraph('March 2000 to July 2000: Programmer at ABM Info Tech (Pvt.) Ltd')
doc.add_paragraph('• Supervise existing applications that was already running in the organization')
doc.add_paragraph('• Analysis, designing and implementation of the financial cycle')

doc.add_paragraph('Sep 1998 – Feb 2000: Application Programmer at ABM Info Sys (Pvt.) Ltd')
doc.add_paragraph('• Generator Set Rental System, Midland Motors Islamabad')
doc.add_paragraph('• Accident Underwriting, Marine Underwriting, Fire Underwriting applications, Habib Insurance Karachi')
doc.add_paragraph('• Call login and inventory control systems, ABM Data Systems and ABM Info Tech')

doc.add_paragraph('Sep 1996 – Feb 1998: Computer Instructor at NEW Computer Guide Institute')
doc.add_paragraph('• Teaching of Information Technology Diploma classes as well as other short & long courses')

# Conferences and Seminars
doc.add_heading('Attended National & International Conferences and Seminars', level=2)
doc.add_paragraph('• 17th Karachi Literary Festival 07th Feb 2026 – Participated as a moderator')
doc.add_paragraph('• Critical Analysis of Ghulam Bahar\'s new book ceremony 10th Dec 2025, presented paper')
doc.add_paragraph('• Balochi Culture at University of Karachi on 26 Sep 2025, Presented paper')
doc.add_paragraph('• Sangat Adabi Diwwan, Karachi, 28th May 2025, Presented Paper on "Impact of Artificial Intelligence on Balochi Literature"')
doc.add_paragraph('• Karachi Literary Festival Feb 2025, Participated as an Urdu Poetess')
doc.add_paragraph('• Ruzhn labzaanki Diwwan, Book launching ceremony at 5th May 2024, participated as a speaker')
doc.add_paragraph('• 16th Aalmi Urdu Conference Karachi, Dec 2023 participated as a panelist')
doc.add_paragraph('• Adab Fest 2022 Karachi, 27th Nov 2022, participated as panelist')
doc.add_paragraph('• International Balochi language Conference at Uppsala University, Sweden from 24th March 2014')

# Personal Information
doc.add_heading('Personal Information', level=2)
doc.add_paragraph('Father Name: Muhammad Suleman Raees')
doc.add_paragraph('GitHub profile: https://github.com/zahidaraees')

# References
doc.add_heading('References', level=2)
doc.add_paragraph('For technical Reference: Muhammad Husain, G.M. at ABM Info Sys, Owner of Diversa Logic')
doc.add_paragraph('Cell: +92-333-211-0907, Email: mhusain1@yahoo.com')
doc.add_paragraph('For Balochi Reference: Dr. Hamid Ali Baloch, Assistant Professor, Balochi Department, University of Balochistan')
doc.add_paragraph('Cell: 03325337818')

doc.save('D:\\resume-bulder\\Resume-with-Unique-Path-and-Shareable-Link\\Zahida_Raees_CV.docx')
print('CV DOCX file created successfully!')
